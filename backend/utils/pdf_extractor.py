"""
Module d'extraction de texte à partir de fichiers PDF et DOCX
avec fonctionnalités d'organisation et d'analyse de contenu.
"""

import os
import re
import tempfile
from datetime import datetime, timedelta
from io import BytesIO
import logging
from typing import List, Dict, Any, Optional, Tuple

# Configuration du logging
logger = logging.getLogger(__name__)

# Tentative d'importation de PyPDF2
try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    logger.warning("PyPDF2 n'est pas installé. L'extraction PDF utilisera pdfminer si disponible.")
    PYPDF2_AVAILABLE = False

# Tentative d'importation de pdfminer.six
try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
    from pdfminer.pdfparser import PDFSyntaxError
    PDFMINER_AVAILABLE = True
except ImportError:
    logger.warning("pdfminer.six n'est pas installé. L'extraction PDF pourrait être limitée.")
    PDFMINER_AVAILABLE = False

# Tentative d'importation de python-docx
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx n'est pas installé. L'extraction DOCX ne sera pas disponible.")
    DOCX_AVAILABLE = False

def extract_text_from_pdf(pdf_data: bytes) -> Optional[str]:
    """
    Extrait le texte d'un fichier PDF.
    
    Args:
        pdf_data: Les données PDF sous forme de bytes
        
    Returns:
        str: Le texte extrait du document PDF
        None: En cas d'erreur
    """
    # Normaliser l'entrée en BytesIO
    pdf_data_io = BytesIO(pdf_data)
    
    extracted_text = ""
    last_error = None
    
    # Essayer d'abord avec PyPDF2 s'il est disponible
    if PYPDF2_AVAILABLE:
        try:
            reader = PdfReader(pdf_data_io)
            texts = []
            
            for page in reader.pages:
                texts.append(page.extract_text())
            
            extracted_text = "\n\n".join([t for t in texts if t])
            
            # Si l'extraction a réussi, retourner le texte
            if extracted_text.strip():
                return extracted_text
            
            # Sinon, essayer avec pdfminer
            logger.info("PyPDF2 n'a pas pu extraire de texte, essai avec pdfminer...")
        except Exception as e:
            last_error = str(e)
            logger.error(f"Erreur lors de l'extraction avec PyPDF2: {e}")
    
    # Essayer avec pdfminer.six s'il est disponible
    if PDFMINER_AVAILABLE:
        try:
            # pdfminer nécessite un fichier, donc nous devons sauvegarder les données temporairement
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                # Revenir au début du BytesIO
                pdf_data_io.seek(0)
                tmp_file.write(pdf_data_io.read())
                tmp_path = tmp_file.name
            
            try:
                extracted_text = pdfminer_extract_text(tmp_path)
            finally:
                # Nettoyer le fichier temporaire
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)
            
            return extracted_text
        
        except PDFSyntaxError as e:
            last_error = f"Erreur de syntaxe PDF: {str(e)}"
            logger.error(f"Erreur de syntaxe PDF: {e}")
        except Exception as e:
            last_error = str(e)
            logger.error(f"Erreur lors de l'extraction avec pdfminer: {e}")
    
    # Si aucune méthode n'a fonctionné, retourner None
    if not extracted_text:
        logger.error("Aucune méthode d'extraction n'a pu extraire du texte du PDF")
        if not last_error:
            last_error = "Impossible d'extraire le texte du PDF."
        return None
    
    return extracted_text

def extract_text_from_docx(docx_data: bytes) -> Optional[str]:
    """
    Extrait le texte d'un fichier DOCX.
    
    Args:
        docx_data: Les données DOCX sous forme de bytes
        
    Returns:
        str: Le texte extrait du document DOCX
        None: En cas d'erreur
    """
    if not DOCX_AVAILABLE:
        logger.error("[DOCX_DEBUG] python-docx n'est pas installé, impossible d'extraire le texte du DOCX")
        
        # Extraction de secours si python-docx n'est pas disponible
        # Pour éviter les erreurs "No /Root object! - Is this really a PDF?" quand on importe un DOCX
        try:
            # Rechercher du texte dans le contenu binaire (méthode de secours)
            text_content = docx_data.decode('utf-8', errors='ignore')
            
            # Nettoyer le texte (enlever les caractères non imprimables)
            import re
            text_content = re.sub(r'[^\x20-\x7E\n\r\t\u00A0-\u00FF\u0100-\u017F]', ' ', text_content)
            
            # Supprimer les longues séquences de caractères répétés souvent présentes dans les fichiers binaires
            text_content = re.sub(r'([^\w\s])\1{3,}', r'\1\1', text_content)
            
            # Essayer d'identifier et d'extraire des paragraphes significatifs
            paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text_content) if len(p.strip()) > 20]
            
            if paragraphs:
                logger.info("[DOCX_DEBUG] Extraction de secours utilisée: texte extrait du contenu binaire")
                return "\n\n".join(paragraphs)
            else:
                logger.warning("[DOCX_DEBUG] Extraction de secours: aucun paragraphe significatif trouvé")
                return "Ce document semble être au format DOCX mais n'a pas pu être analysé correctement.\n\nVeuillez installer python-docx pour une meilleure prise en charge des documents DOCX."
        except Exception as fallback_error:
            logger.error(f"[DOCX_DEBUG] Erreur lors de l'extraction de secours: {str(fallback_error)}")
            return "Impossible d'extraire le contenu de ce document DOCX. Veuillez installer python-docx pour la prise en charge des documents DOCX."
    
    try:
        # Sauvegarder les données dans un fichier temporaire
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            tmp_file.write(docx_data)
            tmp_path = tmp_file.name
            logger.info(f"[DOCX_DEBUG] Fichier temporaire créé: {tmp_path}")
        
        try:
            # Ouvrir le document avec python-docx
            logger.info(f"[DOCX_DEBUG] Tentative d'ouverture du fichier DOCX: {tmp_path}")
            doc = docx.Document(tmp_path)
            logger.info(f"[DOCX_DEBUG] Document DOCX ouvert avec succès: {len(doc.paragraphs)} paragraphes")
            
            # Méthode améliorée pour extraire le texte
            full_text = []
            
            # Extraire le texte de chaque paragraphe
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():  # Ignorer les paragraphes vides
                    full_text.append(para.text.strip())
                    if i < 3:  # Afficher les 3 premiers paragraphes pour débogage
                        logger.info(f"[DOCX_DEBUG] Paragraphe {i+1}: '{para.text[:50]}...' ({len(para.text)} caractères)")
            
            logger.info(f"[DOCX_DEBUG] {len(full_text)} paragraphes extraits du document")
            
            # Extraire le texte des tableaux
            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        full_text.append(" | ".join(row_cells))
            
            logger.info(f"[DOCX_DEBUG] {table_count} tableaux extraits du document")
            
            # Joindre tous les paragraphes avec des sauts de ligne
            extracted_text = "\n".join(full_text)
            
            # Vérification de sécurité pour garantir un contenu minimal
            if not extracted_text or len(extracted_text) < 10:
                logger.warning(f"[DOCX_DEBUG] Le texte extrait est trop court ({len(extracted_text) if extracted_text else 0} caractères)")
                
                # Essayer une méthode alternative avec une extraction paragraphe par paragraphe
                alt_text = []
                for para in doc.paragraphs:
                    para_text = para.text
                    alt_text.append(para_text)
                
                # Essayer également l'extraction via les runs (portions de texte formatées)
                run_text = []
                for para in doc.paragraphs:
                    para_runs = []
                    for run in para.runs:
                        if run.text.strip():
                            para_runs.append(run.text)
                    if para_runs:
                        run_text.append(" ".join(para_runs))
                
                # Comparer les résultats des différentes méthodes
                full_alt_text = "\n".join(alt_text)
                full_run_text = "\n".join(run_text)
                
                logger.info(f"[DOCX_DEBUG] Comparaison des méthodes d'extraction:")
                logger.info(f"[DOCX_DEBUG] - Méthode standard: {len(extracted_text)} caractères")
                logger.info(f"[DOCX_DEBUG] - Méthode alternative (paragraphes): {len(full_alt_text)} caractères")
                logger.info(f"[DOCX_DEBUG] - Méthode alternative (runs): {len(full_run_text)} caractères")
                
                # Choisir la méthode qui produit le texte le plus long
                if len(full_alt_text) > len(extracted_text) and len(full_alt_text) > len(full_run_text):
                    extracted_text = full_alt_text
                    logger.info(f"[DOCX_DEBUG] Méthode alternative (paragraphes) utilisée: {len(extracted_text)} caractères")
                elif len(full_run_text) > len(extracted_text):
                    extracted_text = full_run_text
                    logger.info(f"[DOCX_DEBUG] Méthode alternative (runs) utilisée: {len(extracted_text)} caractères")
            
            # Ajouter le nom du fichier en tant que note
            if extracted_text:
                extracted_text += f"\n\nNote: Ce document a été importé depuis un fichier DOCX."
            
            # Log des informations sur le texte extrait
            logger.info(f"[DOCX_DEBUG] Texte extrait du DOCX: {len(extracted_text)} caractères")
            if extracted_text:
                logger.info(f"[DOCX_DEBUG] Début du texte: {extracted_text[:100]}...")
            
            return extracted_text
        finally:
            # Nettoyer le fichier temporaire
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
                logger.info(f"[DOCX_DEBUG] Fichier temporaire supprimé: {tmp_path}")
    
    except Exception as e:
        logger.error(f"[DOCX_DEBUG] Erreur lors de l'extraction du texte du DOCX: {e}")
        import traceback
        logger.error(f"[DOCX_DEBUG] Traceback: {traceback.format_exc()}")
        
        # Extraction de secours en cas d'erreur
        try:
            # Rechercher du texte dans le contenu binaire
            text_content = docx_data.decode('utf-8', errors='ignore')
            
            # Nettoyer le texte
            import re
            text_content = re.sub(r'[^\x20-\x7E\n\r\t\u00A0-\u00FF\u0100-\u017F]', ' ', text_content)
            text_content = re.sub(r'([^\w\s])\1{3,}', r'\1\1', text_content)
            
            # Extraire des paragraphes significatifs
            paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text_content) if len(p.strip()) > 20]
            
            if paragraphs:
                logger.info("[DOCX_DEBUG] Extraction de secours après erreur: texte extrait du contenu binaire")
                return "\n\n".join(paragraphs)
        except:
            pass
            
        return "Impossible d'extraire le contenu de ce document DOCX suite à une erreur."

def extract_dates_from_text(text: str) -> List[Dict[str, Any]]:
    """
    Recherche et extrait les dates présentes dans le texte avec des informations contextuelles.
    
    Args:
        text: Le texte à analyser
        
    Returns:
        List[Dict]: Liste de dictionnaires contenant:
            - position: position dans le texte
            - date: date au format ISO
            - original: texte original de la date
            - score: score de pertinence (0-100)
            - is_primary: si cette date est susceptible d'être la date principale
            - context: contexte autour de la date
    """
    # Pattern pour les dates au format français et international
    date_patterns = [
        # Format français avec jour de la semaine
        r'(Lundi|Mardi|Mercredi|Jeudi|Vendredi|Samedi|Dimanche)\s+(\d{1,2})\s+(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre)\s+(\d{4})',
        # Format français sans jour de la semaine
        r'(\d{1,2})\s+(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre)\s+(\d{4})',
        # Format ISO
        r'(\d{4})[/-](\d{1,2})[/-](\d{1,2})',
        # Format européen
        r'(\d{1,2})[/-](\d{1,2})[/-](\d{4})',
    ]
    
    # Marqueurs temporels relatifs
    relative_time_patterns = [
        # Aujourd'hui, hier, demain, etc.
        r'\b(aujourd\'hui|hier|avant[\s-]hier|demain|après[\s-]demain)\b',
        # La semaine dernière, le mois prochain, etc.
        r'\b(la\s+semaine\s+dernière|la\s+semaine\s+prochaine|le\s+mois\s+dernier|le\s+mois\s+prochain)\b',
        # Jour de la semaine relatif (lundi dernier, etc.)
        r'\b(lundi|mardi|mercredi|jeudi|vendredi|samedi|dimanche)(\s+dernier|\s+prochain)\b'
    ]
    
    # Dictionnaire de conversion mois français -> numéro
    month_to_number = {
        'janvier': '01', 'février': '02', 'mars': '03', 'avril': '04',
        'mai': '05', 'juin': '06', 'juillet': '07', 'août': '08',
        'septembre': '09', 'octobre': '10', 'novembre': '11', 'décembre': '12'
    }
    
    # Rechercher toutes les occurrences potentielles de dates
    dates_found = []
    
    # Logging pour le debug
    logger.info(f"Texte à analyser pour les dates: {text[:200]}..." if len(text) > 200 else text)
    
    # 1. Extraire les dates absolues
    for pattern in date_patterns:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            date_str = match.group(0)
            position = match.start()
            
            logger.info(f"Date absolue trouvée: {date_str} à la position {position}")
            
            # Extraire le contexte (30 caractères avant et après la date)
            start_context = max(0, position - 30)
            end_context = min(len(text), position + len(date_str) + 30)
            context = text[start_context:end_context]
            
            # Convertir la date au format ISO (YYYY-MM-DD)
            try:
                # Vérifier si c'est une date en français avec un mois en lettres
                has_month_name = False
                for month_name in month_to_number.keys():
                    if month_name in date_str.lower():
                        has_month_name = True
                        break
                
                if has_month_name:
                    # Format français avec ou sans jour de la semaine
                    parts = date_str.split()
                    
                    # Extraire les composants en fonction du format
                    if parts[0].lower() in ['lundi', 'mardi', 'mercredi', 'jeudi', 'vendredi', 'samedi', 'dimanche']:
                        # Avec jour de la semaine
                        day = parts[1]
                        month = parts[2].lower()
                        year = parts[3]
                    else:
                        # Sans jour de la semaine
                        day = parts[0]
                        month = parts[1].lower()
                        year = parts[2]
                    
                    month_num = month_to_number.get(month, '01')
                    iso_date = f"{year}-{month_num}-{day.zfill(2)}"
                elif '-' in date_str or '/' in date_str:
                    # Format ISO ou européen
                    separator = '-' if '-' in date_str else '/'
                    parts = date_str.split(separator)
                    
                    if len(parts[0]) == 4:
                        # Format ISO (YYYY-MM-DD)
                        year = parts[0]
                        month = parts[1].zfill(2)
                        day = parts[2].zfill(2)
                    else:
                        # Format européen (DD-MM-YYYY)
                        day = parts[0].zfill(2)
                        month = parts[1].zfill(2)
                        year = parts[2]
                    
                    iso_date = f"{year}-{month}-{day}"
                else:
                    # Format non reconnu
                    continue
                
                # Vérifier si la date est valide
                datetime.strptime(iso_date, "%Y-%m-%d")
                
                # Analyser le contexte pour déterminer l'importance de cette date
                score, is_primary = analyze_date_context(date_str, context, position)
                
                dates_found.append({
                    "position": position,
                    "date": iso_date,
                    "original": date_str,
                    "score": score,
                    "is_primary": is_primary,
                    "context": context
                })
                
            except ValueError:
                # Date invalide, ignorer
                continue
    
    # 2. Extraire les dates relatives (aujourd'hui, hier, etc.)
    today = datetime.now().date()
    for pattern in relative_time_patterns:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            relative_date_str = match.group(0).lower()
            position = match.start()
            
            logger.info(f"Date relative trouvée: {relative_date_str} à la position {position}")
            
            # Extraire le contexte
            start_context = max(0, position - 30)
            end_context = min(len(text), position + len(relative_date_str) + 30)
            context = text[start_context:end_context]
            
            # Convertir la date relative en date absolue
            relative_date = today  # Par défaut
            
            try:
                if 'aujourd\'hui' in relative_date_str:
                    relative_date = today
                elif 'hier' in relative_date_str:
                    relative_date = today - timedelta(days=1)
                elif 'avant-hier' in relative_date_str:
                    relative_date = today - timedelta(days=2)
                elif 'demain' in relative_date_str:
                    relative_date = today + timedelta(days=1)
                elif 'après-demain' in relative_date_str:
                    relative_date = today + timedelta(days=2)
                elif 'semaine dernière' in relative_date_str:
                    relative_date = today - timedelta(days=7)
                elif 'semaine prochaine' in relative_date_str:
                    relative_date = today + timedelta(days=7)
                elif 'mois dernier' in relative_date_str:
                    # Simplification: 30 jours
                    relative_date = today - timedelta(days=30)
                elif 'mois prochain' in relative_date_str:
                    # Simplification: 30 jours
                    relative_date = today + timedelta(days=30)
                elif any(day in relative_date_str for day in ['lundi', 'mardi', 'mercredi', 'jeudi', 'vendredi', 'samedi', 'dimanche']):
                    # Traitement des jours de la semaine relatifs
                    day_names = {'lundi': 0, 'mardi': 1, 'mercredi': 2, 'jeudi': 3, 'vendredi': 4, 'samedi': 5, 'dimanche': 6}
                    
                    # Trouver quel jour est mentionné
                    day_name = next(day for day in day_names.keys() if day in relative_date_str)
                    target_weekday = day_names[day_name]
                    
                    # Calculer les jours jusqu'au prochain jour de la semaine spécifié
                    days_ahead = target_weekday - today.weekday()
                    if days_ahead <= 0:  # Target day already happened this week
                        days_ahead += 7
                    
                    if 'dernier' in relative_date_str:
                        # La semaine dernière
                        days_ahead -= 14  # Aller deux semaines en arrière puis avancer
                    
                    relative_date = today + timedelta(days=days_ahead)
                
                # Ajouter cette date avec un score élevé (date relative = haute importance)
                iso_date = relative_date.strftime("%Y-%m-%d")
                score, is_primary = analyze_date_context(relative_date_str, context, position, is_relative=True)
                
                dates_found.append({
                    "position": position,
                    "date": iso_date,
                    "original": relative_date_str,
                    "score": score,
                    "is_primary": is_primary,
                    "context": context
                })
                
            except Exception as e:
                logger.error(f"Erreur lors de la conversion de la date relative: {str(e)}")
                continue
    
    # Trier les dates par score puis par position
    dates_found.sort(key=lambda x: (-x["score"], x["position"]))
    
    # Simplification: si plusieurs dates ont un score élevé (> 70), marquer seulement la première comme primaire
    if dates_found and dates_found[0]["score"] > 70:
        dates_found[0]["is_primary"] = True
        for i in range(1, len(dates_found)):
            dates_found[i]["is_primary"] = False
    
    # Compatibilité avec l'ancienne structure
    date_positions = [(date["position"], date["date"]) for date in dates_found]
    date_positions.sort(key=lambda x: x[0])  # Trier par position
    
    logger.info(f"Dates analysées: {dates_found}")
    return date_positions, dates_found

def extract_date_from_filename(filename: str) -> Optional[str]:
    """
    Extrait une date à partir du nom de fichier.
    Supporte les formats comme "Jeudi 19 septembre 2024".
    
    Args:
        filename: Le nom du fichier
        
    Returns:
        str: Date au format ISO (YYYY-MM-DD), ou None si aucune date n'est trouvée
    """
    if not filename:
        logger.warning("[DATE_DEBUG] Extraction de date impossible: aucun nom de fichier fourni")
        return None
    
    # Nettoyer les chemins et les extensions
    base_filename = os.path.basename(filename)
    base_filename = os.path.splitext(base_filename)[0]
    
    logger.info(f"[DATE_DEBUG] Tentative d'extraction de date à partir du nom de fichier: {base_filename}")
    
    # Patterns pour les formats de date dans les noms de fichiers
    patterns = [
        # Format français avec jour de la semaine
        r'(Lundi|Mardi|Mercredi|Jeudi|Vendredi|Samedi|Dimanche)\s+(\d{1,2})\s+(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre)\s+(\d{4})',
        # Format français sans jour de la semaine
        r'(\d{1,2})\s+(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre)\s+(\d{4})',
    ]
    
    # Dictionnaire de conversion mois français -> numéro
    month_to_number = {
        'janvier': '01', 'février': '02', 'mars': '03', 'avril': '04',
        'mai': '05', 'juin': '06', 'juillet': '07', 'août': '08',
        'septembre': '09', 'octobre': '10', 'novembre': '11', 'décembre': '12'
    }
    
    for pattern in patterns:
        # Essayer avec le nom de fichier complet
        match = re.search(pattern, filename, re.IGNORECASE)
        if not match:
            # Essayer avec le nom de fichier de base (sans chemin ni extension)
            match = re.search(pattern, base_filename, re.IGNORECASE)
        
        if match:
            date_str = match.group(0)
            logger.info(f"[DATE_DEBUG] Date trouvée dans le nom de fichier: '{date_str}'")
            parts = date_str.split()
            logger.info(f"[DATE_DEBUG] Parties extraites: {parts}")
            
            # Extraire jour, mois, année selon le format
            if parts[0].lower() in ['lundi', 'mardi', 'mercredi', 'jeudi', 'vendredi', 'samedi', 'dimanche']:
                day = parts[1]
                month = parts[2].lower()
                year = parts[3]
                logger.info(f"[DATE_DEBUG] Format avec jour de semaine détecté: jour={day}, mois={month}, année={year}")
            else:
                day = parts[0]
                month = parts[1].lower()
                year = parts[2]
                logger.info(f"[DATE_DEBUG] Format sans jour de semaine détecté: jour={day}, mois={month}, année={year}")
            
            # Convertir au format ISO
            month_num = month_to_number.get(month, '01')
            iso_date = f"{year}-{month_num}-{day.zfill(2)}"
            
            try:
                # Vérifier si la date est valide
                datetime.strptime(iso_date, "%Y-%m-%d")
                logger.info(f"[DATE_DEBUG] Date extraite du nom de fichier et validée: {iso_date}")
                return iso_date
            except ValueError as e:
                logger.warning(f"[DATE_DEBUG] Date invalide extraite du nom de fichier: {iso_date}. Erreur: {str(e)}")
                continue
    
    logger.warning("[DATE_DEBUG] Aucune date trouvée dans le nom de fichier")
    return None

def analyze_date_context(date_str: str, context: str, position: int, is_relative: bool = False) -> Tuple[int, bool]:
    """
    Analyse le contexte autour d'une date pour déterminer son importance.
    
    Args:
        date_str: Texte de la date
        context: Contexte autour de la date
        position: Position de la date dans le texte
        is_relative: Si la date est relative (aujourd'hui, hier, etc.)
        
    Returns:
        Tuple[int, bool]: (score de pertinence (0-100), si c'est probablement la date principale)
    """
    score = 50  # Score par défaut
    
    # Facteurs qui augmentent l'importance
    importance_indicators = [
        (r'^[^\.]*' + re.escape(date_str), 25),  # Date au début d'une phrase
        (r'^' + re.escape(date_str), 40),  # Date au début du texte
        (r'le ' + re.escape(date_str), 10),  # "Le" avant la date
        (r'date.*?:.*?' + re.escape(date_str), 30),  # Format "Date: ..."
        (r'jour.*?:.*?' + re.escape(date_str), 30),  # Format "Jour: ..."
        (r'\. ' + re.escape(date_str), 20),  # Date juste après un point (début de phrase)
    ]
    
    # Facteurs qui diminuent l'importance (références au passé/futur)
    diminish_indicators = [
        (r'depuis (le|la|les) ' + re.escape(date_str), -20),  # "depuis le..."
        (r'avant (le|la|les) ' + re.escape(date_str), -20),  # "avant le..."
        (r'après (le|la|les) ' + re.escape(date_str), -20),  # "après le..."
        (r'jusqu\'(au|à la) ' + re.escape(date_str), -15),  # "jusqu'au..."
        (r'à partir (du|de la) ' + re.escape(date_str), -15),  # "à partir du..."
        (r'commencé (le|la|en) ' + re.escape(date_str), -10),  # "commencé le..."
        (r'terminé (le|la|en) ' + re.escape(date_str), -10),  # "terminé le..."
    ]
    
    # Analyser les indicateurs
    for pattern, points in importance_indicators:
        if re.search(pattern, context, re.IGNORECASE):
            score += points
    
    for pattern, points in diminish_indicators:
        if re.search(pattern, context, re.IGNORECASE):
            score += points
    
    # Les dates relatives ont généralement plus d'importance
    if is_relative:
        score += 15
    
    # Les dates en début de document ont plus d'importance
    if position < 200:
        score += 20
    
    # Limiter le score entre 0 et 100
    score = max(0, min(100, score))
    
    # Déterminer si c'est probablement la date principale
    is_primary = score > 70
    
    return score, is_primary

def extract_automatic_tags(text: str, threshold: float = 0.01) -> List[str]:
    """
    Extrait automatiquement des tags à partir du texte en privilégiant les termes techniques et sujets pertinents.
    
    Args:
        text: Le texte à analyser
        threshold: Seuil de fréquence pour considérer un mot comme tag
        
    Returns:
        List[str]: Liste de tags potentiels
    """
    # Liste de sujets techniques pertinents à rechercher prioritairement
    technical_subjects = [
        "microsoft", "sharepoint", "azure", "aws", "google", "cloud", "devops", 
        "kubernetes", "docker", "python", "javascript", "typescript", "react", "angular", 
        "vue", "nodejs", "database", "sql", "nosql", "mongodb", "postgresql", "mysql",
        "api", "rest", "graphql", "microservices", "backend", "frontend", "fullstack",
        "agile", "scrum", "kanban", "jira", "git", "github", "gitlab", "cicd", "jenkins",
        "terraform", "ansible", "cybersecurity", "machine learning", "intelligence artificielle",
        "ia", "data science", "big data", "hadoop", "spark", "etl", "kafka", "elasticsearch",
        "web", "mobile", "app", "application", "testing", "automation", "integration",
        "php", "java", "spring", "dotnet", "csharp", "c#", "interface", "architecture",
        "powerbi", "power automate", "power apps", "flow", "automate", "workflow", "dashboard",
        "rapport", "projet", "étude", "développement", "programmation", "application", "formation",
        "équipe", "réunion", "daily", "meeting", "présentation", "documentation", "rapport",
        "client", "ticketing", "résolution", "bug", "problème", "solution", "déploiement",
        "technique", "technologie", "innovation", "digital", "numérique", "optimisation"
    ]
    
    # Liste de mots à ne jamais inclure comme tags car ils sont liés à l'import et pas au contenu
    blacklisted_tags = [
        "import", "erreur", "importerreur", "error", "date_from_filename",
        "fichier", "document", "extraction", "texte", "contenu", "analyse"
    ]
    
    # Extraction des mots (sans ponctuation, chiffres, etc.)
    words = re.findall(r'\b[a-zA-ZÀ-ÿ]{4,}\b', text.lower())
    
    # Liste étendue de mots vides français pour un filtrage plus efficace
    stopwords = set([
        'dans', 'avec', 'pour', 'cette', 'mais', 'avoir', 'faire', 'plus', 'tout', 'bien', 
        'être', 'comme', 'nous', 'leur', 'sans', 'vous', 'dont', 'alors', 'aussi', 'donc', 
        'cela', 'ceux', 'celle', 'celui', 'entre', 'pendant', 'depuis', 'notre', 'votre', 
        'avons', 'sommes', 'sont', 'était', 'étaient', 'sera', 'seront', 'avez', 'ainsi',
        'ceci', 'cela', 'celle', 'celui', 'ceux', 'chaque', 'comment', 'contre', 'dans', 
        'depuis', 'devant', 'donc', 'elle', 'elles', 'encore', 'entre', 'vers', 'voici', 
        'voilà', 'votre', 'vouloir', 'vous', 'aucun', 'auquel', 'autre', 'autres', 'assez',
        'après', 'avant', 'avoir', 'beaucoup', 'chaque', 'chez', 'comment', 'dedans', 'dehors',
        'déjà', 'dessous', 'dessus', 'donc', 'durant', 'enfin', 'encore', 'ensuite', 'être',
        'falloir', 'faire', 'haut', 'jusqu', 'lequel', 'lorsque', 'maintenant', 'moins', 'même',
        'plupart', 'plusieurs', 'pouvoir', 'presque', 'puis', 'puisque', 'quand', 'quelque',
        'savoir', 'sinon', 'tandis', 'tellement', 'toujours', 'toutefois', 'trop', 'très',
        'cette', 'cela', 'celui', 'celle', 'ceux', 'celles', 'autre', 'autres', 'même',
        'notamment', 'également', 'ainsi', 'donc', 'afin', 'après', 'avant', 'pendant', 'selon',
        'malgré', 'suite', 'voici', 'voilà', 'certes', 'ensuite', 'bref', 'concernant', 'concerné'
    ])
    
    # Ajouter les mots blacklistés aux stopwords
    stopwords.update(blacklisted_tags)
    
    # Filtrer les mots vides
    words = [w for w in words if w not in stopwords]
    
    # Compter les occurrences
    word_counts = {}
    for word in words:
        word_counts[word] = word_counts.get(word, 0) + 1
    
    total_words = len(words)
    if total_words == 0:
        # S'il n'y a pas de mots significatifs, chercher spécifiquement les sujets techniques
        # Pour éviter de retourner des tags comme "import" ou "erreur"
        for subject in technical_subjects:
            if subject in text.lower():
                return [subject]
        # Si vraiment rien n'est trouvé, retourner un tag générique pertinent
        return ["projet"]
    
    # Rechercher des sujets techniques connus en priorité
    technical_tags = []
    for subject in technical_subjects:
        if subject in text.lower() and subject not in technical_tags:
            technical_tags.append(subject)
            # Retirer les occurrences de ce sujet pour éviter les doublons
            if subject in word_counts:
                del word_counts[subject]
    
    # Vérifier si "import" ou "erreur" sont présents dans les mots comptés et les supprimer
    for blacklisted in blacklisted_tags:
        if blacklisted in word_counts:
            del word_counts[blacklisted]
    
    # Sélectionner les autres mots significatifs qui dépassent le seuil
    other_tags = [word for word, count in word_counts.items() 
                  if count / total_words > threshold and count > 1]
    
    # Combiner les tags techniques et les autres tags significatifs
    combined_tags = technical_tags + other_tags
    
    # Vérification finale pour s'assurer qu'aucun tag blacklisté n'est présent
    combined_tags = [tag for tag in combined_tags if tag not in blacklisted_tags]
    
    # Si aucun tag n'a été trouvé après tout le filtrage, utiliser un tag par défaut
    if not combined_tags:
        return ["projet"]
    
    # Limiter le nombre de tags (en privilégiant les tags techniques)
    return combined_tags[:5]

def analyze_entry_content(text: str) -> Dict[str, Any]:
    """
    Analyse le contenu d'une entrée pour en extraire des informations.
    
    Args:
        text: Le texte à analyser
        
    Returns:
        Dict: Informations extraites (type, entreprise, tags)
    """
    result = {
        "type_entree": "quotidien",  # valeur par défaut
        "entreprise_id": None,
        "tags": extract_automatic_tags(text)
    }
    
    # Détection de type d'entrée
    if re.search(r'\b(formation|cours|apprendre|étudier|apprentissage)\b', text.lower()):
        result["type_entree"] = "formation"
    elif re.search(r'\b(projet|développement|application|implémentation|feature)\b', text.lower()):
        result["type_entree"] = "projet"
    elif re.search(r'\b(réflexion|analyse|pensée|considération|bilan)\b', text.lower()):
        result["type_entree"] = "réflexion"
    
    return result

def extract_date_from_filename(filename: str) -> Optional[str]:
    """
    Extrait une date à partir du nom de fichier.
    Supporte les formats comme "Jeudi 19 septembre 2024".
    
    Args:
        filename: Le nom du fichier
        
    Returns:
        str: Date au format ISO (YYYY-MM-DD), ou None si aucune date n'est trouvée
    """
    if not filename:
        logger.warning("[DATE_DEBUG] Extraction de date impossible: aucun nom de fichier fourni")
        return None
    
    # Nettoyer les chemins et les extensions
    base_filename = os.path.basename(filename)
    base_filename = os.path.splitext(base_filename)[0]
    
    logger.info(f"[DATE_DEBUG] Tentative d'extraction de date à partir du nom de fichier: {base_filename}")
    
    # Patterns pour les formats de date dans les noms de fichiers
    patterns = [
        # Format français avec jour de la semaine
        r'(Lundi|Mardi|Mercredi|Jeudi|Vendredi|Samedi|Dimanche)\s+(\d{1,2})\s+(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre)\s+(\d{4})',
        # Format français sans jour de la semaine
        r'(\d{1,2})\s+(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre)\s+(\d{4})',
    ]
    
    # Dictionnaire de conversion mois français -> numéro
    month_to_number = {
        'janvier': '01', 'février': '02', 'mars': '03', 'avril': '04',
        'mai': '05', 'juin': '06', 'juillet': '07', 'août': '08',
        'septembre': '09', 'octobre': '10', 'novembre': '11', 'décembre': '12'
    }
    
    for pattern in patterns:
        # Essayer avec le nom de fichier complet
        match = re.search(pattern, filename, re.IGNORECASE)
        if not match:
            # Essayer avec le nom de fichier de base (sans chemin ni extension)
            match = re.search(pattern, base_filename, re.IGNORECASE)
        
        if match:
            date_str = match.group(0)
            logger.info(f"[DATE_DEBUG] Date trouvée dans le nom de fichier: '{date_str}'")
            parts = date_str.split()
            logger.info(f"[DATE_DEBUG] Parties extraites: {parts}")
            
            # Extraire jour, mois, année selon le format
            if parts[0].lower() in ['lundi', 'mardi', 'mercredi', 'jeudi', 'vendredi', 'samedi', 'dimanche']:
                day = parts[1]
                month = parts[2].lower()
                year = parts[3]
                logger.info(f"[DATE_DEBUG] Format avec jour de semaine détecté: jour={day}, mois={month}, année={year}")
            else:
                day = parts[0]
                month = parts[1].lower()
                year = parts[2]
                logger.info(f"[DATE_DEBUG] Format sans jour de semaine détecté: jour={day}, mois={month}, année={year}")
            
            # Convertir au format ISO
            month_num = month_to_number.get(month, '01')
            iso_date = f"{year}-{month_num}-{day.zfill(2)}"
            
            try:
                # Vérifier si la date est valide
                datetime.strptime(iso_date, "%Y-%m-%d")
                logger.info(f"[DATE_DEBUG] Date extraite du nom de fichier et validée: {iso_date}")
                return iso_date
            except ValueError as e:
                logger.warning(f"[DATE_DEBUG] Date invalide extraite du nom de fichier: {iso_date}. Erreur: {str(e)}")
                continue
    
    logger.warning("[DATE_DEBUG] Aucune date trouvée dans le nom de fichier")
    return None

def process_document(file_content: bytes, filename: Optional[str] = None) -> List[Dict[str, Any]]:
    """
    Traite un document (PDF ou DOCX) et extrait son contenu sous forme d'entrées de journal.
    
    Args:
        file_content: Le contenu du fichier
        filename: Le nom du fichier
        
    Returns:
        List[Dict]: Liste des entrées extraites
    """
    # Vérifier d'abord si le contenu du fichier et le nom du fichier sont bien définis
    if not file_content:
        logger.error("Aucun contenu de fichier fourni")
        return []
        
    if not filename:
        logger.warning("Aucun nom de fichier fourni, tentative de traitement malgré tout")
        
    # Extraire la date du nom du fichier en priorité
    filename_date = None
    if filename:
        logger.info(f"[DATE_DEBUG] Tentative d'extraction de date à partir du nom du fichier: {filename}")
        filename_date = extract_date_from_filename(filename)
        if filename_date:
            logger.info(f"[DATE_DEBUG] Date extraite avec succès du nom du fichier: {filename_date}")
            logger.warning(f"[DATE_PRIORITY] Date du nom de fichier sera utilisée: {filename_date}")
        else:
            logger.warning(f"[DATE_PRIORITY] Aucune date trouvée dans le nom du fichier: {filename}")
    
    # Créer une entrée artificielle en cas d'échec complet
    # Essayons d'extraire du texte même sans contenu significatif
    sample_text = "Ce document a été importé mais son contenu n'a pas pu être analysé correctement. " + \
                 "Veuillez vérifier que vous avez installé les bibliothèques nécessaires (python-docx pour les fichiers DOCX, PyPDF2 pour les PDF)."
    
    default_entry = [{
        "date": filename_date if filename_date else datetime.now().strftime("%Y-%m-%d"),
        "texte": f"Document {filename} importé le {datetime.now().strftime('%d/%m/%Y à %H:%M')}. \n\n{sample_text}",
        "type_entree": "quotidien",
        "tags": ["projet", "document"],  # Utiliser des tags génériques et utiles
        "source_document": filename,
        "date_source": "filename" if filename_date else "current"
    }]
    
    text = None
    file_text_extracted = False
    
    # Détection du type de fichier par l'extension
    is_docx = False
    is_pdf = False
    
    if filename:
        filename_lower = filename.lower()
        if filename_lower.endswith('.docx') or 'docx' in filename_lower:
            is_docx = True
            logger.info(f"[FILE_DEBUG] Fichier détecté comme DOCX par le nom: {filename}")
        elif filename_lower.endswith('.pdf') or 'pdf' in filename_lower:
            is_pdf = True
            logger.info(f"[FILE_DEBUG] Fichier détecté comme PDF par le nom: {filename}")
    
    # Si c'est un fichier DOCX, utiliser une extraction spécifique DOCX en priorité
    if is_docx:
        logger.info(f"[FILE_DEBUG] Traitement prioritaire comme fichier DOCX: {filename}")
        try:
            # Extraction par python-docx
            text = extract_text_from_docx(file_content)
            if text and len(text) > 50:  # Vérifier qu'on a extrait du contenu significatif
                logger.info(f"[FILE_DEBUG] Texte extrait avec succès du fichier DOCX: {len(text)} caractères")
                file_text_extracted = True
            else:
                logger.warning(f"[FILE_DEBUG] Extraction DOCX a retourné peu de contenu ({len(text) if text else 0} caractères)")
                
                # Tentative d'extraction par méthode directe (même si python-docx n'est pas disponible)
                try:
                    # Méthode d'urgence: extraire directement du contenu binaire
                    raw_text = file_content.decode('utf-8', errors='ignore')
                    # Nettoyer et extraire du texte lisible
                    import re
                    # Supprimer les caractères binaires et garder l'alphanumérique et la ponctuation
                    clean_text = re.sub(r'[^\x20-\x7E\n\r\t\u00A0-\u00FF\u0100-\u017F]', ' ', raw_text)
                    # Extraire des blocs de texte significatifs
                    blocks = re.findall(r'[A-Za-z0-9\s.,;:!?«»()\[\]\'\"]{20,}', clean_text)
                    if blocks:
                        # Concaténer les blocs trouvés
                        fallback_text = "\n\n".join(blocks)
                        if len(fallback_text) > 100:
                            text = fallback_text
                            logger.info(f"[FILE_DEBUG] Texte extrait par méthode de secours: {len(text)} caractères")
                            file_text_extracted = True
                except Exception as raw_err:
                    logger.error(f"[FILE_DEBUG] Erreur lors de l'extraction brute: {str(raw_err)}")
        except Exception as e:
            logger.error(f"[FILE_DEBUG] Erreur lors de l'extraction du fichier DOCX: {str(e)}")
            text = None
    
    # Si c'est un fichier PDF ou si le traitement DOCX a échoué
    if (is_pdf or not file_text_extracted) and not is_docx:
        try:
            text = extract_text_from_pdf(file_content)
            if text and len(text) > 50:
                logger.info(f"[FILE_DEBUG] Texte extrait avec succès du fichier PDF: {len(text)} caractères")
                file_text_extracted = True
            else:
                logger.warning(f"[FILE_DEBUG] Extraction PDF a retourné peu de contenu ({len(text) if text else 0} caractères)")
        except Exception as e:
            logger.error(f"[FILE_DEBUG] Erreur lors de l'extraction du fichier PDF: {str(e)}")
            text = None
    
    # Tentative de secours: extraction directe de texte à partir des données binaires
    if not file_text_extracted:
        try:
            logger.info("[FILE_DEBUG] Tentative d'extraction de texte par recherche directe dans les données binaires")
            # Convertir les données binaires en texte
            raw_text = file_content.decode('utf-8', errors='ignore')
            
            # Nettoyer et trouver des blocs de texte exploitables
            import re
            # Supprimer les caractères binaires/non-imprimables
            clean_text = re.sub(r'[^\x20-\x7E\n\r\t\u00A0-\u00FF\u0100-\u017F]', ' ', raw_text)
            # Réduire les espaces multiples
            clean_text = re.sub(r'\s+', ' ', clean_text)
            
            # Rechercher des phrases ou paragraphes significatifs
            # (recherche de séquences de mots et ponctuation d'au moins 30 caractères)
            text_blocks = re.findall(r'[A-Za-z0-9àáâäæçèéêëìíîïòóôöùúûüÿœÀÁÂÄÆÇÈÉÊËÌÍÎÏÒÓÔÖÙÚÛÜŸŒ\s.,;:!?«»()\[\]\'\"]{30,}', clean_text)
            
            if text_blocks:
                # Filtrer pour garder uniquement les blocs avec un ratio raisonnable de lettres/chiffres
                filtered_blocks = []
                for block in text_blocks:
                    # Compter les caractères alphanumériques
                    alnum_count = sum(1 for c in block if c.isalnum())
                    # Si au moins 40% du bloc est composé de caractères alphanumériques
                    if alnum_count / len(block) > 0.4:
                        filtered_blocks.append(block)
                
                if filtered_blocks:
                    extracted_text = "\n\n".join(filtered_blocks)
                    if len(extracted_text) > 100:  # Si on a au moins 100 caractères
                        text = extracted_text
                        logger.info(f"[FILE_DEBUG] Texte extrait par analyse directe: {len(text)} caractères")
                        file_text_extracted = True
        except Exception as fallback_error:
            logger.error(f"[FILE_DEBUG] Erreur lors de l'extraction de secours: {str(fallback_error)}")
    
    # Si aucune méthode n'a fonctionné, utiliser l'entrée par défaut
    if not file_text_extracted or not text:
        logger.error(f"[FILE_DEBUG] Toutes les méthodes d'extraction ont échoué pour: {filename}")
        return default_entry
    
    if not text:
        logger.error(f"Impossible d'extraire du texte du document '{filename}'")
        return []
    
    # Vérifier que le texte extrait est valide et contient un minimum de contenu
    if len(text.strip()) < 10:
        # Si le texte est trop court, créer une entrée artificielle pour éviter l'échec
        logger.warning(f"Texte extrait trop court ({len(text.strip())} caractères), création d'une entrée artificielle")
        # Utiliser la date du nom de fichier si disponible, sinon la date actuelle
        entry_date = filename_date if filename_date else datetime.now().strftime("%Y-%m-%d")
        date_source = "filename" if filename_date else "current"
        logger.info(f"[DATE_DEBUG] Entrée artificielle créée avec date {entry_date} (source: {date_source})")
        return [{
            "date": entry_date,
            "texte": f"Document {filename} importé le {datetime.now().strftime('%d/%m/%Y à %H:%M')}. \n\nLe contenu n'a pas pu être extrait correctement.",
            "type_entree": "quotidien",
            "tags": ["import", "erreur"],
            "source_document": filename,
            "date_source": date_source
        }]
    
    # Rechercher les dates pour diviser le contenu
    date_positions, dates_analyzed = extract_dates_from_text(text)
    
    # Logging pour debug
    logger.info(f"Dates trouvées dans le document: {len(dates_analyzed)} dates")
    
    # Extraire la date du nom de fichier en priorité
    filename_date = None
    if filename:
        filename_date = extract_date_from_filename(filename)
        if filename_date:
            logger.info(f"[DATE_DEBUG] Date extraite du nom de fichier: {filename_date}")
            
            # IMPORTANT: Utiliser directement la date du nom de fichier et ignorer toute autre date
            # Créer directement l'entrée avec cette date
            metadata = analyze_entry_content(text)
            logger.info(f"[DATE_DEBUG] Création d'une entrée avec la date du fichier: {filename_date}")
            logger.warning(f"[DATE_PRIORITY] Utilisation de la date extraite du nom de fichier pour l'importation: {filename_date}")
            
            # Utiliser uniquement les tags significatifs extraits du contenu
            tags = metadata["tags"]
            
            return [{
                "date": filename_date,
                "texte": text,
                "type_entree": metadata["type_entree"],
                "tags": tags,
                "source_document": filename,
                "date_source": "filename"  # Marqueur pour indiquer la source de la date
            }]
    
    # Si aucune date n'est trouvée dans le nom du fichier, continuer avec la recherche dans le contenu
    logger.info(f"[DATE_DEBUG] Aucune date dans le nom de fichier, recherche dans le contenu")
    
    # Si aucune date n'est trouvée dans le contenu non plus
    if not dates_analyzed:
        current_date = datetime.now().strftime("%Y-%m-%d")
        logger.info(f"[DATE_DEBUG] Aucune date trouvée, utilisation de la date actuelle: {current_date}")
        
        metadata = analyze_entry_content(text)
        return [{
            "date": current_date,
            "texte": text,
            "type_entree": metadata["type_entree"],
            "tags": metadata["tags"],
            "source_document": filename,
            "date_source": "current"  # Marqueur pour indiquer la source de la date
        }]
    
    # Stratégie de traitement des dates:
    # 1. Trouver la date principale (avec le meilleur score)
    # 2. Créer une entrée pour cette date avec le texte complet
    # 3. Si d'autres dates sont pertinentes, créer des entrées séparées
    
    # Ce bloc est désormais ignoré car la date est extraite et utilisée en amont
    # Nous gardons cette logique pour la compatibilité, mais elle ne devrait pas être utilisée
    
    # Identifier la date principale à partir du contenu
    primary_dates = [d for d in dates_analyzed if d["is_primary"]]
    if primary_dates:
        primary_date = primary_dates[0]
    else:
        # Si aucune date n'est marquée comme principale, prendre celle avec le score le plus élevé
        primary_date = dates_analyzed[0] if dates_analyzed else {"date": datetime.now().strftime("%Y-%m-%d"), "score": 0, "is_primary": False, "position": 0, "context": ""}
    
    logger.info(f"[DATE_DEBUG] Date principale identifiée dans le contenu: {primary_date['date']} (score: {primary_date['score']})")
    logger.info(f"[DATE_DEBUG] Note: Cette date sera ignorée car nous utilisons déjà la date du fichier en priorité.")
    
    # Créer une entrée principale avec la date principale et tout le texte
    metadata = analyze_entry_content(text)
    primary_entry = {
        "date": primary_date["date"],
        "texte": text,
        "type_entree": metadata["type_entree"],
        "tags": metadata["tags"],
        "source_document": filename,
        "primary_date": True  # Marquer cette entrée comme utilisant la date principale
    }
    
    # Décider si nous devons créer des entrées séparées pour d'autres dates
    # Stratégie: ne créer des entrées séparées que pour les dates avec un score élevé
    # et qui ne sont pas trop proches de la date principale
    
    secondary_entries = []
    date_diff_threshold = 7  # Différence minimale en jours
    score_threshold = 65    # Score minimum pour les dates secondaires
    
    for date_info in dates_analyzed:
        # Ne pas traiter à nouveau la date principale
        if date_info["date"] == primary_date["date"]:
            continue
        
        # Vérifier si cette date a un score assez élevé
        if date_info["score"] < score_threshold:
            continue
        
        # Calculer la différence en jours entre cette date et la date principale
        try:
            this_date = datetime.strptime(date_info["date"], "%Y-%m-%d").date()
            primary_d = datetime.strptime(primary_date["date"], "%Y-%m-%d").date()
            days_diff = abs((this_date - primary_d).days)
            
            if days_diff < date_diff_threshold:
                # Dates trop proches, ignorer
                continue
                
            # Extraire un contexte significatif autour de cette date
            pos = date_info["position"]
            max_context_length = 1000  # Caractères
            
            # Trouver où commencer la découpe (paragraphe ou phrase)
            context_start = max(0, pos - 100)
            # Reculer jusqu'au début d'une phrase ou d'un paragraphe
            while context_start > 0 and text[context_start] not in ".!?\n":
                context_start -= 1
            if context_start > 0:
                context_start += 1  # Avancer après le marqueur de fin
            
            # Trouver où terminer la découpe
            context_end = min(len(text), pos + 900)
            # Avancer jusqu'à la fin d'une phrase ou d'un paragraphe
            while context_end < len(text) - 1 and text[context_end] not in ".!?\n":
                context_end += 1
            if context_end < len(text) - 1:
                context_end += 1  # Inclure le marqueur de fin
            
            # Extraire le contexte
            context_text = text[context_start:context_end].strip()
            
            # Vérifier si le contexte est suffisamment long
            if len(context_text) >= 10:
                # Créer une entrée secondaire
                secondary_metadata = analyze_entry_content(context_text)
                secondary_entry = {
                    "date": date_info["date"],
                    "texte": context_text,
                    "type_entree": secondary_metadata["type_entree"],
                    "tags": secondary_metadata["tags"],
                    "source_document": filename,
                    "primary_date": False  # Marquer cette entrée comme utilisant une date secondaire
                }
                secondary_entries.append(secondary_entry)
                logger.info(f"Entrée secondaire créée avec date {date_info['date']} (score: {date_info['score']}, contexte: {len(context_text)} caractères)")
        except Exception as e:
            logger.error(f"Erreur lors de la création d'une entrée secondaire: {str(e)}")
    
    # Compiler les entrées (principale + secondaires) selon la stratégie choisie
    entries = [primary_entry]
    
    # Option: ajouter des entrées secondaires (désactivé par défaut)
    include_secondary_entries = False  # Paramètre configurable
    
    if include_secondary_entries and secondary_entries:
        entries.extend(secondary_entries)
    
    return entries

# Garder la fonction existante pour la compatibilité
def process_pdf_file(file_content: bytes, filename: Optional[str] = None) -> List[Dict[str, Any]]:
    """
    Traite un fichier PDF et extrait son contenu sous forme d'entrées de journal.
    
    Args:
        file_content: Le contenu du fichier PDF
        filename: Le nom du fichier
        
    Returns:
        List[Dict]: Liste des entrées extraites
    """
    return process_document(file_content, filename)