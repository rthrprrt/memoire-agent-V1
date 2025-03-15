import io
import json
from typing import List, Dict
import logging

from pydantic import BaseModel
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)

# Modèle pour les options d'export
class ExportOptions(BaseModel):
    format: str = "pdf"  # "pdf" ou "docx"
    include_toc: bool = True
    include_bibliography: bool = True
    include_appendices: bool = True
    page_numbers: bool = True
    cover_page: bool = True
    document_title: str = "Mémoire de Mission Professionnelle"
    author_name: str = ""
    institution_name: str = "Epitech Digital School"
    academic_year: str = "2024-2025"
    margin_top_cm: float = 2.5
    margin_bottom_cm: float = 2.5
    margin_left_cm: float = 3.0
    margin_right_cm: float = 2.5

class MemoryExporter:
    """
    Classe pour exporter le mémoire dans différents formats
    avec mise en page académique.
    """
    def __init__(self, memory_manager):
        self.memory_manager = memory_manager

    async def export_to_pdf(self, options: ExportOptions) -> bytes:
        """Exporte le mémoire au format PDF"""
        # Récupérer les sections et le plan
        outline = await self.memory_manager.get_outline()
        sections_data = await self._gather_sections_content(outline)
        
        # Récupérer la bibliographie si nécessaire
        bibliography = []
        if options.include_bibliography:
            bibliography = await self._gather_bibliography()
        
        # Création du document PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            topMargin=options.margin_top_cm * 28.35,
            bottomMargin=options.margin_bottom_cm * 28.35,
            leftMargin=options.margin_left_cm * 28.35,
            rightMargin=options.margin_right_cm * 28.35
        )
        
        # Définition des styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='Heading1',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            spaceBefore=24
        ))
        styles.add(ParagraphStyle(
            name='Heading2',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=10,
            spaceBefore=18
        ))
        styles.add(ParagraphStyle(
            name='Heading3',
            parent=styles['Heading3'],
            fontSize=12,
            spaceAfter=8,
            spaceBefore=12
        ))
        styles.add(ParagraphStyle(
            name='Normal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=10,
            leading=14
        ))
        
        content = []
        
        # Page de couverture
        if options.cover_page:
            content.append(Spacer(1, 100))
            content.append(Paragraph(options.document_title, ParagraphStyle(
                name='Title', parent=styles['Title'], fontSize=24, alignment=1
            )))
            content.append(Spacer(1, 50))
            if options.author_name:
                content.append(Paragraph(f"Par: {options.author_name}", ParagraphStyle(
                    name='Author', parent=styles['Normal'], fontSize=14, alignment=1
                )))
            content.append(Spacer(1, 20))
            content.append(Paragraph(options.institution_name, ParagraphStyle(
                name='Institution', parent=styles['Normal'], fontSize=14, alignment=1
            )))
            content.append(Paragraph(options.academic_year, ParagraphStyle(
                name='Year', parent=styles['Normal'], fontSize=14, alignment=1
            )))
            content.append(PageBreak())
        
        # Table des matières
        if options.include_toc:
            content.append(Paragraph("Table des matières", styles['Heading1']))
            for section in sections_data:
                level = section.get("level", 0)
                indent = "    " * level
                content.append(Paragraph(f"{indent}{section['title']}", styles['Normal']))
            content.append(PageBreak())
        
        # Contenu des sections
        for section in sections_data:
            level = section.get("level", 0)
            if level == 0:
                content.append(Paragraph(section['title'], styles['Heading1']))
            elif level == 1:
                content.append(Paragraph(section['title'], styles['Heading2']))
            else:
                content.append(Paragraph(section['title'], styles['Heading3']))
            
            if section['content']:
                paragraphs = section['content'].split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        content.append(Paragraph(para, styles['Normal']))
            content.append(Spacer(1, 12))
        
        # Bibliographie
        if options.include_bibliography and bibliography:
            content.append(PageBreak())
            content.append(Paragraph("Bibliographie", styles['Heading1']))
            for ref in bibliography:
                content.append(Paragraph(f"• {ref['citation']}", ParagraphStyle(
                    name='Bibliography', parent=styles['Normal'], 
                    leftIndent=36, firstLineIndent=-36
                )))
                content.append(Spacer(1, 6))
        
        doc.build(content)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        return pdf_bytes

    async def export_to_docx(self, options: ExportOptions) -> bytes:
        """Exporte le mémoire au format DOCX (Word)"""
        outline = await self.memory_manager.get_outline()
        sections_data = await self._gather_sections_content(outline)
        bibliography = []
        if options.include_bibliography:
            bibliography = await self._gather_bibliography()
        
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(options.margin_top_cm)
            section.bottom_margin = Cm(options.margin_bottom_cm)
            section.left_margin = Cm(options.margin_left_cm)
            section.right_margin = Cm(options.margin_right_cm)
        
        styles = doc.styles
        
        # Style pour les titres de niveau 1
        h1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.base_style = styles['Heading 1']
        h1_style.font.size = Pt(16)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0, 0, 0)
        
        # Style pour les titres de niveau 2
        h2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.base_style = styles['Heading 2']
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
        
        # Style pour les titres de niveau 3
        h3_style = styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
        h3_style.base_style = styles['Heading 3']
        h3_style.font.size = Pt(12)
        h3_style.font.bold = True
        
        # Style pour le texte normal
        normal_style = styles['Normal']
        normal_style.font.size = Pt(11)
        
        # Page de couverture
        if options.cover_page:
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(options.document_title)
            title_run.font.size = Pt(24)
            title_run.font.bold = True
            doc.add_paragraph()
            if options.author_name:
                author_para = doc.add_paragraph()
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                author_run = author_para.add_run(f"Par: {options.author_name}")
                author_run.font.size = Pt(14)
            doc.add_paragraph()
            institution_para = doc.add_paragraph()
            institution_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            inst_run = institution_para.add_run(options.institution_name)
            inst_run.font.size = Pt(14)
            year_para = doc.add_paragraph()
            year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            year_run = year_para.add_run(options.academic_year)
            year_run.font.size = Pt(14)
            doc.add_page_break()
        
        # Table des matières (champ Word)
        if options.include_toc:
            doc.add_heading("Table des matières", level=1)
            p = doc.add_paragraph()
            run = p.add_run("\tCLICK TO UPDATE TABLE OF CONTENTS")
            run.font.color.rgb = RGBColor(128, 128, 128)
            p = doc.add_paragraph()
            run = p.add_run("{TOC \\o \"1-3\" \\h \\z \\u}")
            run.font.color.rgb = RGBColor(128, 128, 128)
            doc.add_page_break()
        
        # Contenu des sections
        for section in sections_data:
            doc.add_heading(section['title'], level=section.get("level", 0) + 1)
            if section['content']:
                paragraphs = section['content'].split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        p = doc.add_paragraph()
                        p.add_run(para.strip())
        
        # Bibliographie
        if options.include_bibliography and bibliography:
            doc.add_page_break()
            doc.add_heading("Bibliographie", level=1)
            for ref in bibliography:
                doc.add_paragraph(ref['citation'], style='List Bullet')
        
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        return docx_bytes

    async def _gather_sections_content(self, outline: List[Dict], level: int = 0) -> List[Dict]:
        """Récupère le contenu de toutes les sections de manière récursive"""
        result = []
        for section in outline:
            section_data = await self.memory_manager.get_section(section["id"])
            result.append({
                "id": section["id"],
                "title": section_data.get("titre", ""),
                "content": section_data.get("contenu", ""),
                "level": level
            })
            if "children" in section and section["children"]:
                children_data = await self._gather_sections_content(section["children"], level + 1)
                result.extend(children_data)
        return result

    async def _gather_bibliography(self) -> List[Dict]:
        """Récupère toutes les références bibliographiques"""
        conn = self.memory_manager.get_connection()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT * FROM bibliography_references ORDER BY authors, year
        """)
        results = cursor.fetchall()
        conn.close()
        references = []
        for result in results:
            # Conversion des auteurs (en supposant qu'ils soient stockés au format JSON)
            authors_raw = json.loads(result["authors"]) if isinstance(result["authors"], str) else result["authors"]
            authors = ", ".join(authors_raw)
            if len(authors_raw) > 3:
                authors = f"{authors_raw[0]} et al."
            year = result["year"]
            title = result["title"]
            publisher = result.get("publisher") or ""
            citation = f"{authors} ({year}). {title}. {publisher}."
            references.append({
                "id": result["id"],
                "type": result["type"],
                "citation": citation
            })
        return references
