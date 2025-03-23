import os
import io
import uuid
from datetime import datetime
from typing import Dict, Any, Optional, List
import logging

from pydantic import BaseModel

# Importer les modules d'export conditionnellement
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, Cm
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class ExportOptions(BaseModel):
    """Options pour l'export de document"""
    format: str = "pdf"  # "pdf" ou "docx"
    include_toc: bool = True
    include_bibliography: bool = True
    include_appendices: bool = True
    page_numbers: bool = True
    cover_page: bool = True
    document_title: str = "Mémoire de Mission Professionnelle"
    author_name: str = ""
    institution_name: str = "Epitech Digital School"
    academic_year: str = "2024-2025"
    margin_top_cm: float = 2.5
    margin_bottom_cm: float = 2.5
    margin_left_cm: float = 3.0
    margin_right_cm: float = 2.5

class ExportService:
    """Service pour gérer l'export de documents"""
    
    def __init__(self, export_dir: str):
        self.export_dir = export_dir
        os.makedirs(export_dir, exist_ok=True)
        
        # Vérifier les dépendances disponibles
        if not REPORTLAB_AVAILABLE:
            logger.warning("ReportLab n'est pas installé. L'export PDF sera limité.")
        if not DOCX_AVAILABLE:
            logger.warning("python-docx n'est pas installé. L'export DOCX sera limité.")
    
    async def export_to_pdf(self, content: Dict[str, Any], options: ExportOptions) -> bytes:
        """Exporte le contenu au format PDF"""
        if not REPORTLAB_AVAILABLE:
            return b"%PDF-1.4\n1 0 obj\n<< /Title (PDF Export Unavailable) >>\nendobj\ntrailer\n<< /Root 1 0 R >>\n%%EOF"
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            topMargin=options.margin_top_cm * 28.35,
            bottomMargin=options.margin_bottom_cm * 28.35,
            leftMargin=options.margin_left_cm * 28.35,
            rightMargin=options.margin_right_cm * 28.35
        )
        
        styles = getSampleStyleSheet()
        flowables = []
        
        # Page de couverture
        if options.cover_page:
            flowables.append(Spacer(1, 100))
            flowables.append(Paragraph(options.document_title, styles['Title']))
            flowables.append(Spacer(1, 50))
            
            if options.author_name:
                flowables.append(Paragraph(f"Par: {options.author_name}", styles['Normal']))
                flowables.append(Spacer(1, 20))
            
            flowables.append(Paragraph(options.institution_name, styles['Normal']))
            flowables.append(Paragraph(options.academic_year, styles['Normal']))
            flowables.append(PageBreak())
        
        # Table des matières
        if options.include_toc:
            flowables.append(Paragraph("Table des matières", styles['Heading1']))
            flowables.append(Spacer(1, 20))
            
            if 'sections' in content:
                # Générer la table des matières
                for section in content['sections']:
                    level = section.get('level', 0)
                    indent = "    " * level
                    flowables.append(Paragraph(f"{indent}{section.get('title', 'Sans titre')}", styles['Normal']))
            
            flowables.append(PageBreak())
        
        # Contenu des sections
        if 'sections' in content:
            for section in content['sections']:
                level = section.get('level', 0)
                title = section.get('title', 'Sans titre')
                content_text = section.get('content', '')
                
                # Ajouter le titre avec le style approprié
                if level == 0:
                    flowables.append(Paragraph(title, styles['Heading1']))
                elif level == 1:
                    flowables.append(Paragraph(title, styles['Heading2']))
                else:
                    flowables.append(Paragraph(title, styles['Heading3']))
                
                flowables.append(Spacer(1, 10))
                
                # Découper le contenu en paragraphes
                if content_text:
                    paragraphs = content_text.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            flowables.append(Paragraph(para.strip(), styles['Normal']))
                            flowables.append(Spacer(1, 5))
        
        # Bibliographie
        if options.include_bibliography and 'bibliography' in content:
            flowables.append(PageBreak())
            flowables.append(Paragraph("Bibliographie", styles['Heading1']))
            flowables.append(Spacer(1, 20))
            
            for ref in content.get('bibliography', []):
                citation = ref.get('citation', '')
                if citation:
                    flowables.append(Paragraph(f"• {citation}", styles['Normal']))
                    flowables.append(Spacer(1, 5))
        
        # Construire le document
        doc.build(flowables)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        return pdf_bytes
    
    async def export_to_docx(self, content: Dict[str, Any], options: ExportOptions) -> bytes:
        """Exporte le contenu au format DOCX"""
        if not DOCX_AVAILABLE:
            # Retourner un fichier DOCX minimal
            return b'PK\x03\x04\x14\x00\x00\x00\x00\x00\x00\x00!\x00\x00\x00\x00\x00\x00\x00\x00\x00'
        
        doc = Document()
        
        # Configurer les marges
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(options.margin_top_cm)
            section.bottom_margin = Cm(options.margin_bottom_cm)
            section.left_margin = Cm(options.margin_left_cm)
            section.right_margin = Cm(options.margin_right_cm)
        
        # Page de couverture
        if options.cover_page:
            doc.add_paragraph().add_run(options.document_title).bold = True
            
            if options.author_name:
                doc.add_paragraph(f"Par: {options.author_name}")
            
            doc.add_paragraph(options.institution_name)
            doc.add_paragraph(options.academic_year)
            doc.add_page_break()
        
        # Table des matières
        if options.include_toc:
            doc.add_heading("Table des matières", level=1)
            
            if 'sections' in content:
                # Générer la table des matières
                for section in content['sections']:
                    level = section.get('level', 0)
                    indent = "    " * level
                    p = doc.add_paragraph(indent)
                    p.add_run(section.get('title', 'Sans titre'))
            
            doc.add_page_break()
        
        # Contenu des sections
        if 'sections' in content:
            for section in content['sections']:
                level = section.get('level', 0)
                title = section.get('title', 'Sans titre')
                content_text = section.get('content', '')
                
                # Ajouter le titre avec le niveau approprié
                doc.add_heading(title, level=level+1)
                
                # Découper le contenu en paragraphes
                if content_text:
                    paragraphs = content_text.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
        
        # Bibliographie
        if options.include_bibliography and 'bibliography' in content:
            doc.add_page_break()
            doc.add_heading("Bibliographie", level=1)
            
            for ref in content.get('bibliography', []):
                citation = ref.get('citation', '')
                if citation:
                    p = doc.add_paragraph("• ")
                    p.add_run(citation)
        
        # Enregistrer le document
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        
        return docx_bytes
    
    async def save_document(self, document_bytes: bytes, format: str, title: str) -> Dict[str, Any]:
        """Sauvegarde un document exporté"""
        document_id = str(uuid.uuid4())
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        sanitized_title = "".join(c if c.isalnum() or c in "._- " else "_" for c in title)
        filename = f"{sanitized_title}_{timestamp}.{format}"
        
        file_path = os.path.join(self.export_dir, filename)
        
        # Sauvegarder le document
        with open(file_path, 'wb') as f:
            f.write(document_bytes)
        
        document_info = {
            "id": document_id,
            "title": title,
            "format": format,
            "filename": filename,
            "created_at": datetime.now().isoformat(),
            "file_path": file_path,
            "file_size": len(document_bytes)
        }
        
        # Sauvegarder les métadonnées
        metadata_path = os.path.join(self.export_dir, f"{document_id}_meta.json")
        with open(metadata_path, 'w') as f:
            import json
            json.dump(document_info, f, indent=2)
        
        return document_info
    
    async def get_document(self, document_id: str) -> Optional[Dict[str, Any]]:
        """Récupère un document exporté"""
        metadata_path = os.path.join(self.export_dir, f"{document_id}_meta.json")
        
        if not os.path.exists(metadata_path):
            return None
        
        with open(metadata_path, 'r') as f:
            import json
            document_info = json.load(f)
        
        if not os.path.exists(document_info.get('file_path', '')):
            return None
        
        with open(document_info['file_path'], 'rb') as f:
            document_info['content'] = f.read()
        
        return document_info
    
    async def list_documents(self, limit: int = 20) -> List[Dict[str, Any]]:
        """Liste les documents exportés"""
        import glob
        import json
        
        metadata_files = glob.glob(os.path.join(self.export_dir, "*_meta.json"))
        documents = []
        
        for metadata_path in metadata_files[:limit]:
            try:
                with open(metadata_path, 'r') as f:
                    document_info = json.load(f)
                
                # Ne pas inclure le contenu du document
                if 'content' in document_info:
                    del document_info['content']
                
                # Vérifier si le fichier existe toujours
                document_info['available'] = os.path.exists(document_info.get('file_path', ''))
                
                documents.append(document_info)
            except Exception as e:
                logger.error(f"Erreur lors de la lecture des métadonnées {metadata_path}: {str(e)}")
        
        # Trier par date de création (les plus récents en premier)
        documents.sort(key=lambda x: x.get('created_at', ''), reverse=True)
        
        return documents

# Export service singleton
_export_service = None

def get_export_service():
    """Obtient l'instance singleton du service d'export"""
    global _export_service
    if _export_service is None:
        from core.config import settings
        export_dir = settings.EXPORT_PATH
        _export_service = ExportService(export_dir)
    return _export_service

async def create_export(
    memory_manager, 
    export_service: ExportService,
    options: ExportOptions
) -> Dict[str, Any]:
    """
    Crée un export du mémoire
    
    Args:
        memory_manager: Gestionnaire de mémoire
        export_service: Service d'export
        options: Options d'export
        
    Returns:
        Dict: Informations sur le document exporté
    """
    # Récupérer les données pour l'export
    outline = await memory_manager.get_outline()
    
    # Récupérer le contenu des sections
    sections = []
    
    # Fonction récursive pour aplatir l'outline
    async def flatten_outline(items, level=0):
        for item in items:
            # Récupérer le contenu de la section
            section = await memory_manager.get_memoire_section(item['id'])
            if section:
                sections.append({
                    'id': section['id'],
                    'title': section['titre'],
                    'content': section.get('contenu', ''),
                    'level': level
                })
            
            # Traiter les enfants récursivement
            if 'children' in item and item['children']:
                await flatten_outline(item['children'], level + 1)
    
    # Aplatir l'outline pour récupérer toutes les sections
    await flatten_outline(outline)
    
    # Récupérer la bibliographie
    bibliography = []  # Implémenter la récupération de la bibliographie
    
    # Préparer le contenu complet
    content = {
        'sections': sections,
        'bibliography': bibliography
    }
    
    # Générer le document selon le format
    if options.format.lower() == 'pdf':
        document_bytes = await export_service.export_to_pdf(content, options)
    elif options.format.lower() == 'docx':
        document_bytes = await export_service.export_to_docx(content, options)
    else:
        raise ValueError(f"Format non supporté: {options.format}")
    
    # Sauvegarder le document
    document_info = await export_service.save_document(
        document_bytes=document_bytes,
        format=options.format.lower(),
        title=options.document_title
    )
    
    return document_info